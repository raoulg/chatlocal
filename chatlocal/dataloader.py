from __future__ import annotations

from pathlib import Path
from typing import Iterator, List

from docx import Document as WordDocument
import pdftotext
from loguru import logger
import json

from chatlocal.settings import Document, FileType, ParserSettings


class DataStore:
    def __init__(self) -> None:
        self.data: List[Document] = []
        self._index: int = 0

    def add(self, document: Document) -> None:
        self.data.append(document)

    def __len__(self) -> int:
        return len(self.data)

    def __getitem__(self, index: int) -> Document:
        return self.data[index]

    def __iter__(self) -> Iterator[Document]:  # type: ignore
        self._index = 0
        return self

    def __repr__(self) -> str:
        return f"Datastore({len(self.data)} Documents))"

    def __next__(self) -> Document:
        if self._index < len(self.data):
            result = self.data[self._index]
            self._index += 1
            return result
        else:
            raise StopIteration

    class Config:
        arbitrary_types_allowed = True


class Parser:
    def __init__(self, settings: ParserSettings) -> None:
        self._settings = settings

    @property
    def settings(self) -> ParserSettings:
        return self._settings

    def __call__(self, source: Path, ftype: FileType) -> Document:
        if ftype in self.settings.textfiles:
            return self.parse_text(source)
        if ftype in self.settings.wordfiles:
            return self.parse_word(source)
        if ftype in self.settings.pdffiles:
            return self.parse_pdf(source)
        if ftype in self.settings.jupyterfiles:
            return self.parse_jupyter(source)

        raise ValueError(f"Filetype {type} not supported")

    def parse_text(self, source: Path) -> Document:
        with open(source, "r") as f:
            content = f.read()
        document = Document(text=content, source=source)
        return document

    def parse_word(self, source: Path) -> Document:
        worddoc = WordDocument(source)
        content = ""
        for par in worddoc.paragraphs:
            content += par.text + "\n"
        document = Document(text=content, source=source)
        return document

    def parse_pdf(self, source: Path) -> Document:
        try:
            with open(source, "rb") as f:
                pdf = pdftotext.PDF(f)
            content = "\n\n".join(pdf)
            document = Document(text=content, source=source)
            return document
        except Exception as e:
            logger.error(f"Error parsing PDF {source}: {e}")

    def parse_jupyter(self, source: Path) -> Document:
        with open(source, "r") as f:
            raw = json.load(f)
        content_: List[str] = ["".join(cell['source']) for cell in raw['cells'] if 'source' in cell]
        # join the list of strings into a single string
        content = "".join(content_)

        document = Document(text=content, source=source)
        return document


class DataLoader:
    def __init__(self, filetypes: List[FileType] = [FileType.MD]) -> None:
        logger.info(f"Initializing Dataloader for files of type {filetypes}...")
        self.filetypes = filetypes
        self.datastore = DataStore()
        self.parser = Parser(ParserSettings.from_filetypes(filetypes))

    def load_files(self, path: Path) -> None:
        if path.exists() is False:
            raise FileNotFoundError(f"{path} does not exist")
        if path.is_dir() is False:
            raise NotADirectoryError(f"{path} is not a directory")

        notepaths = self.walk_dir(path)
        num_docs = 0
        skipped = set()
        num_skipped = 0
        for source in notepaths:
            file_extension = source.suffix.lower()  # Get the file extension

            if not any(file_extension == ft.value for ft in self.filetypes):
                skipped.add(file_extension)
                num_skipped += 1
                continue

            document = self.parser(source, FileType(file_extension))
            if len(document.text) > 0:
                self.datastore.add(document)
                num_docs += 1
        logger.info(f"Added {num_docs} notes to the datastore.")
        logger.info(f"Skipped {num_skipped} files with extensions {skipped}.")

    def get_datastore(self) -> DataStore:
        return self.datastore

    def walk_dir(self, path: Path) -> Iterator[Path]:
        """loops recursively through a folder

        Args:
            path (Path): folder to loop trough. If a directory
                is encountered, loop through that recursively.

        Yields:
            Iterator: all paths in a folder and subdirs.
        """
        ignore_dirs = self.parser.settings.ignore_dirs
        path = path.expanduser().resolve()
        for p in Path(path).iterdir():
            if p.is_dir() and p.name in ignore_dirs:
                logger.debug(f"Skipping directory {p}")
                continue
            if p.is_dir():
                yield from self.walk_dir(p)
                continue
            yield p
